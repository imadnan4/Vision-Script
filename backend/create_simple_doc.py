from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_guide():
    doc = Document()
    
    # Title
    title = doc.add_heading('Using Anaconda Libraries Outside Anaconda Prompt', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Quick Guide to Run Python Backend with Anaconda Libraries', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Key Steps
    doc.add_heading('1. Key Steps We Used:', 2)
    steps = [
        'We didn\'t need to reinstall any libraries',
        'We used the Python from Anaconda directly',
        'We fixed the OpenMP error with environment variables'
    ]
    for step in steps:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(step)
    
    # 2. Working Command
    doc.add_heading('2. The Working Command:', 2)
    command = doc.add_paragraph()
    command.style = 'Code'
    command.add_run('$env:KMP_DUPLICATE_LIB_OK="TRUE"\n$env:KMP_WARNINGS="0"\n& "$env:USERPROFILE\\Anaconda3\\python.exe" main_test.py')
    
    # 3. Why This Works
    doc.add_heading('3. Why This Works:', 2)
    reasons = [
        'Uses Anaconda\'s Python directly (from Anaconda3 folder)',
        'Accesses all libraries already installed in Anaconda',
        'Fixes the OpenMP error that was blocking execution',
        'No need to activate conda environment'
    ]
    for reason in reasons:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(reason)
    
    # 4. Error Handling
    doc.add_heading('4. If You Get Errors:', 2)
    errors = [
        'Make sure you\'re using PowerShell',
        'Check if Anaconda is installed in the default location (Users\\YourUsername\\Anaconda3)',
        'If Anaconda is installed elsewhere, replace the path accordingly'
    ]
    for error in errors:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(error)
    
    # 5. Alternative Method
    doc.add_heading('5. Alternative Method:', 2)
    doc.add_paragraph('If the above doesn\'t work, you can also:')
    
    doc.add_paragraph('a. Find your Anaconda Python path:')
    paths = [
        'Usually in C:\\Users\\YourUsername\\Anaconda3\\python.exe',
        'Or wherever you installed Anaconda'
    ]
    for path in paths:
        p = doc.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(path)
    
    doc.add_paragraph('b. Use that full path to run your script:')
    command = doc.add_paragraph()
    command.style = 'Code'
    command.add_run('"C:\\Users\\YourUsername\\Anaconda3\\python.exe" main_test.py')
    
    # Remember note
    note = doc.add_paragraph()
    note.add_run('Remember: ').bold = True
    note.add_run('The key is using Anaconda\'s Python directly with the environment variables set, rather than trying to activate the conda environment.')
    
    # Save the document
    doc.save('Using_Anaconda_Libraries.docx')

if __name__ == '__main__':
    create_guide() 