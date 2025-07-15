import cv2
import easyocr
import pytesseract
import numpy as np
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import tempfile
import os
import base64
import xlsxwriter
from PIL import Image, ImageEnhance
import re
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from sumy.summarizers.luhn import LuhnSummarizer
from sumy.summarizers.text_rank import TextRankSummarizer
import nltk
from collections import Counter
import math
import PyPDF2
import io
import requests
import json
from datetime import datetime, timedelta
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()



app = Flask(__name__)
CORS(app)

# Dictionary to cache EasyOCR readers for different languages
readers = {}

# Default reader (English)
readers['en'] = easyocr.Reader(['en'])

# OpenRouter rate limiting
class OpenRouterRateLimit:
    def __init__(self):
        self.requests_per_minute = 20  # Conservative limit for free tier
        self.requests_per_day = 200    # Daily limit for free tier
        self.request_times = []
        self.daily_requests = []
        self.last_reset = datetime.now()
        self.is_rate_limited = False
        self.rate_limit_until = None

    def reset_if_needed(self):
        now = datetime.now()
        if now - self.last_reset >= timedelta(days=1):
            self.daily_requests = []
            self.last_reset = now
            self.is_rate_limited = False
            self.rate_limit_until = None

    def can_make_request(self):
        self.reset_if_needed()
        now = datetime.now()

        # Check if we're still in cooldown
        if self.rate_limit_until and now < self.rate_limit_until:
            return False, f"Rate limited. Cooldown until {self.rate_limit_until.strftime('%H:%M:%S')}"

        # Remove requests older than 1 minute
        minute_ago = now - timedelta(minutes=1)
        self.request_times = [t for t in self.request_times if t > minute_ago]

        # Remove requests older than 1 day
        day_ago = now - timedelta(days=1)
        self.daily_requests = [t for t in self.daily_requests if t > day_ago]

        # Check minute limit
        if len(self.request_times) >= self.requests_per_minute:
            self.is_rate_limited = True
            self.rate_limit_until = now + timedelta(minutes=1)
            return False, f"Rate limit reached ({self.requests_per_minute}/minute). Try again in 1 minute."

        # Check daily limit
        if len(self.daily_requests) >= self.requests_per_day:
            self.is_rate_limited = True
            self.rate_limit_until = now + timedelta(hours=24)
            return False, f"Daily limit reached ({self.requests_per_day}/day). Try again tomorrow."

        return True, "OK"

    def record_request(self):
        now = datetime.now()
        self.request_times.append(now)
        self.daily_requests.append(now)

    def get_status(self):
        self.reset_if_needed()
        now = datetime.now()
        minute_ago = now - timedelta(minutes=1)
        day_ago = now - timedelta(days=1)

        # Clean old data
        self.request_times = [t for t in self.request_times if t > minute_ago]
        self.daily_requests = [t for t in self.daily_requests if t > day_ago]

        minute_requests = len(self.request_times)
        daily_requests = len(self.daily_requests)

        cooldown_remaining = 0
        if self.rate_limit_until and now < self.rate_limit_until:
            cooldown_remaining = int((self.rate_limit_until - now).total_seconds())

        return {
            "minute_requests": minute_requests,
            "minute_limit": self.requests_per_minute,
            "daily_requests": daily_requests,
            "daily_limit": self.requests_per_day,
            "is_rate_limited": self.is_rate_limited and cooldown_remaining > 0,
            "cooldown_seconds": cooldown_remaining
        }

# Global rate limiter instance
openrouter_rate_limiter = OpenRouterRateLimit()

# Supported languages with their codes
SUPPORTED_LANGUAGES = {
    'en': 'English',
    'fr': 'French',
    'es': 'Spanish',
    'de': 'German',
    'it': 'Italian',
    'pt': 'Portuguese',
    'nl': 'Dutch',
    'zh': 'Chinese',
    'ja': 'Japanese',
    'ko': 'Korean',
    'ru': 'Russian',
    'ar': 'Arabic',
    'hi': 'Hindi'
}

def get_reader(lang_code='en'):
    """Get or create an EasyOCR reader for the specified language"""
    if lang_code not in SUPPORTED_LANGUAGES:
        lang_code = 'en'  # Default to English if unsupported
    
    if lang_code not in readers:
        readers[lang_code] = easyocr.Reader([lang_code])
    
    return readers[lang_code]

def extract_text(file_path, model='easyocr', lang_code='en'):
    if model == 'pytesseract':
        img = Image.open(file_path)
        # For pytesseract, we use the language parameter if available
        if lang_code != 'en':
            return pytesseract.image_to_string(img, lang=lang_code)
        else:
            return pytesseract.image_to_string(img)
    else:
        reader = get_reader(lang_code)
        results = reader.readtext(file_path)
        return ' '.join([res[1] for res in results])

@app.route('/supported_languages', methods=['GET'])
def supported_languages():
    """Return a list of supported languages"""
    return jsonify({
        "languages": SUPPORTED_LANGUAGES
    })

@app.route('/summarization_options', methods=['GET'])
def summarization_options():
    """Return available summarization options"""
    return jsonify({
        "algorithms": {
            "textrank": "TextRank Algorithm",
            "lsa": "Latent Semantic Analysis",
            "luhn": "Luhn Algorithm",
            "abstractive": "Abstractive Summary"
        },
        "smart_options": {
            "local_smart": "Smart Summary (Local)",
            "openrouter": "AI Summary (OpenRouter Mistral)"
        },
        "types": {
            "paragraph": "Paragraph Summary",
            "bullets": "Bullet Points",
            "keyphrases": "Key Phrases"
        },
        "lengths": {
            "short": "Short (25%)",
            "medium": "Medium (50%)",
            "long": "Long (75%)"
        }
    })

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""

        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"

        return text.strip()
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return None

@app.route('/extract_pdf_text', methods=['POST'])
def extract_pdf_text():
    """Extract text from uploaded PDF file"""
    try:
        if 'pdf_file' not in request.files:
            return jsonify({"error": "No PDF file provided"}), 400

        pdf_file = request.files['pdf_file']

        if pdf_file.filename == '':
            return jsonify({"error": "No file selected"}), 400

        if not pdf_file.filename.lower().endswith('.pdf'):
            return jsonify({"error": "File must be a PDF"}), 400

        # Extract text from PDF
        extracted_text = extract_text_from_pdf(pdf_file)

        if extracted_text is None:
            return jsonify({"error": "Failed to extract text from PDF"}), 500

        if not extracted_text.strip():
            return jsonify({"error": "No text found in PDF"}), 400

        return jsonify({
            "text": extracted_text,
            "filename": pdf_file.filename,
            "word_count": len(extracted_text.split()),
            "status": "success"
        })

    except Exception as e:
        return jsonify({
            "error": str(e),
            "status": "error"
        }), 500

@app.route('/upload_image', methods=['POST'])
def upload_image():
    file = request.files['image']
    model = request.form.get('model', 'easyocr').lower()
    lang_code = request.form.get('language', 'en').lower()
    
    file_path = os.path.join(tempfile.gettempdir(), file.filename)
    file.save(file_path)
    
    extracted_text = extract_text(file_path, model, lang_code)
    
    return jsonify({
        "recognized_text": extracted_text
    })

@app.route('/download_format', methods=['POST'])
def download_format():
    chosen_format = request.form.get('format')
    text_data = request.form.get('text_data', '')
    original_text = request.form.get('original_text', '')
    statistics = request.form.get('statistics', '')
    is_summary = request.form.get('is_summary', 'false').lower() == 'true'

    filename_prefix = "summary" if is_summary else "extracted_text"

    if chosen_format == 'txt':
        tmp_path = os.path.join(tempfile.gettempdir(), f"{filename_prefix}.txt")
        with open(tmp_path, 'w', encoding='utf-8') as f:
            if is_summary and original_text:
                f.write("=== ORIGINAL TEXT ===\n\n")
                f.write(original_text)
                f.write("\n\n=== SUMMARY ===\n\n")
                f.write(text_data)
                if statistics:
                    f.write(f"\n\n=== STATISTICS ===\n\n{statistics}")
            else:
                f.write(text_data)
        return send_file(tmp_path, as_attachment=True)

    elif chosen_format == 'docx':
        from docx import Document
        doc = Document()

        if is_summary and original_text:
            doc.add_heading('Text Summary Report', 0)
            doc.add_heading('Original Text', level=1)
            doc.add_paragraph(original_text)
            doc.add_heading('Summary', level=1)
            doc.add_paragraph(text_data)
            if statistics:
                doc.add_heading('Statistics', level=1)
                doc.add_paragraph(statistics)
        else:
            doc.add_paragraph(text_data)

        tmp_path = os.path.join(tempfile.gettempdir(), f"{filename_prefix}.docx")
        doc.save(tmp_path)
        return send_file(tmp_path, as_attachment=True)

    elif chosen_format == 'excel':
        tmp_path = os.path.join(tempfile.gettempdir(), f"{filename_prefix}.xlsx")
        workbook = xlsxwriter.Workbook(tmp_path)
        worksheet = workbook.add_worksheet()

        # Add headers
        header_format = workbook.add_format({'bold': True, 'bg_color': '#333333', 'font_color': 'white'})

        if is_summary and original_text:
            worksheet.write(0, 0, 'Section', header_format)
            worksheet.write(0, 1, 'Content', header_format)

            row = 1
            worksheet.write(row, 0, 'Original Text')
            worksheet.write(row, 1, original_text)
            row += 1

            worksheet.write(row, 0, 'Summary')
            worksheet.write(row, 1, text_data)
            row += 1

            if statistics:
                worksheet.write(row, 0, 'Statistics')
                worksheet.write(row, 1, statistics)
        else:
            worksheet.write(0, 0, 'Content', header_format)
            for idx, line in enumerate(text_data.split('\n')):
                worksheet.write(idx + 1, 0, line)

        workbook.close()
        return send_file(tmp_path, as_attachment=True)

    else:
        return jsonify({"error": "Invalid format"}), 400

def clean_extracted_text(text, lang_code='en'):
    # Default keys (English)
    keys = [
        "Name", "Father Name", "Gender", "Country of Stay", "Identity Number",
        "Date of Birth", "Date of Issue", "Date of Expiry"
    ]
    
    # Language-specific keys
    if lang_code == 'es':  # Spanish
        keys = [
            "Nombre", "Nombre del padre", "Género", "País de residencia", "Número de identidad",
            "Fecha de nacimiento", "Fecha de emisión", "Fecha de caducidad"
        ]
    elif lang_code == 'fr':  # French
        keys = [
            "Nom", "Nom du père", "Sexe", "Pays de séjour", "Numéro d'identité",
            "Date de naissance", "Date d'émission", "Date d'expiration"
        ]
    elif lang_code == 'de':  # German
        keys = [
            "Name", "Vatersname", "Geschlecht", "Aufenthaltsland", "Identitätsnummer",
            "Geburtsdatum", "Ausstellungsdatum", "Ablaufdatum"
        ]
    
    extracted_data = {}
    
    for key in keys:
        if key in text:
            try:
                # Try to find the text between this key and the next key
                next_key_index = keys.index(key) + 1
                if next_key_index < len(keys) and keys[next_key_index] in text:
                    extracted_data[key] = text.split(key)[1].split(keys[next_key_index])[0].strip()
                else:
                    # If there's no next key in the text, take everything after this key
                    extracted_data[key] = text.split(key)[1].strip().split('\n')[0]
            except (IndexError, ValueError):
                # Fallback: just take the first line after the key
                parts = text.split(key)
                if len(parts) > 1:
                    extracted_data[key] = parts[1].strip().split('\n')[0]
    
    return extracted_data

# Initialize NLTK data (download required data if not present)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

def clean_text(text):
    """Clean and preprocess text for summarization"""
    # Remove extra whitespace and normalize
    text = re.sub(r'\s+', ' ', text.strip())
    # Remove special characters but keep basic punctuation
    text = re.sub(r'[^\w\s.,!?;:-]', '', text)
    return text

def get_text_statistics(original_text, summary_text):
    """Calculate statistics for the summarization"""
    original_words = len(original_text.split())
    summary_words = len(summary_text.split())

    reduction_percentage = ((original_words - summary_words) / original_words) * 100 if original_words > 0 else 0

    # Estimate reading time (average 200 words per minute)
    original_reading_time = math.ceil(original_words / 200)
    summary_reading_time = math.ceil(summary_words / 200)

    return {
        'original_word_count': original_words,
        'summary_word_count': summary_words,
        'reduction_percentage': round(reduction_percentage, 1),
        'original_reading_time_minutes': original_reading_time,
        'summary_reading_time_minutes': summary_reading_time
    }

def score_sentences(sentences, text):
    """Score sentences based on multiple factors for better summarization"""
    import string
    from collections import defaultdict

    # Calculate word frequencies
    words = nltk.word_tokenize(text.lower())
    stop_words = set(nltk.corpus.stopwords.words('english'))
    words = [word for word in words if word.isalpha() and word not in stop_words]
    word_freq = Counter(words)

    sentence_scores = defaultdict(float)

    for i, sentence in enumerate(sentences):
        sentence_words = nltk.word_tokenize(sentence.lower())
        sentence_words = [word for word in sentence_words if word.isalpha() and word not in stop_words]

        if len(sentence_words) == 0:
            continue

        # Factor 1: Word frequency score
        freq_score = sum(word_freq.get(word, 0) for word in sentence_words) / len(sentence_words)

        # Factor 2: Position score (earlier sentences are more important)
        position_score = 1.0 - (i / len(sentences)) * 0.5

        # Factor 3: Length score (prefer medium-length sentences)
        length_score = min(len(sentence_words) / 20, 1.0) if len(sentence_words) > 5 else 0.5

        # Factor 4: Keyword density (sentences with more unique important words)
        unique_words = len(set(sentence_words))
        keyword_score = unique_words / len(sentence_words) if len(sentence_words) > 0 else 0

        # Factor 5: Avoid very short or very long sentences
        if len(sentence_words) < 5 or len(sentence_words) > 50:
            length_penalty = 0.5
        else:
            length_penalty = 1.0

        # Combine all factors
        total_score = (freq_score * 0.4 + position_score * 0.2 + length_score * 0.2 + keyword_score * 0.2) * length_penalty
        sentence_scores[i] = total_score

    return sentence_scores

def extractive_summarize(text, algorithm='textrank', sentences_count=3):
    """Perform intelligent extractive summarization"""
    try:
        cleaned_text = clean_text(text)
        if len(cleaned_text.split()) < 20:
            return text  # Return original if too short

        sentences = nltk.sent_tokenize(cleaned_text)
        if len(sentences) <= sentences_count:
            return cleaned_text

        if algorithm == 'smart':
            # Use our custom smart scoring
            sentence_scores = score_sentences(sentences, cleaned_text)

            # Get top sentences by score
            top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:sentences_count]

            # Sort by original order to maintain flow
            top_sentences = sorted(top_sentences, key=lambda x: x[0])

            summary_sentences = [sentences[i] for i, _ in top_sentences]
            return ' '.join(summary_sentences)
        else:
            # Use existing algorithms but with better sentence selection
            parser = PlaintextParser.from_string(cleaned_text, Tokenizer("english"))

            if algorithm == 'lsa':
                summarizer = LsaSummarizer()
            elif algorithm == 'luhn':
                summarizer = LuhnSummarizer()
            else:  # textrank
                summarizer = TextRankSummarizer()

            summarizer.stop_words = nltk.corpus.stopwords.words('english')

            # Get more sentences than needed, then filter
            extended_count = min(sentences_count * 2, len(sentences))
            summary = summarizer(parser.document, extended_count)

            # Score the summary sentences and pick the best ones
            summary_sentences = [str(sentence) for sentence in summary]
            if len(summary_sentences) > sentences_count:
                sentence_scores = score_sentences(summary_sentences, cleaned_text)
                top_indices = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:sentences_count]
                top_indices = sorted(top_indices, key=lambda x: x[0])  # Maintain order
                summary_sentences = [summary_sentences[i] for i, _ in top_indices]

            return ' '.join(summary_sentences)

    except Exception as e:
        print(f"Error in extractive summarization: {e}")
        return text

def bullet_point_summarize(text, max_points=5):
    """Create smart bullet point summary"""
    try:
        cleaned_text = clean_text(text)
        sentences = nltk.sent_tokenize(cleaned_text)

        if len(sentences) <= max_points:
            return '\n'.join([f"• {sentence.strip()}" for sentence in sentences])

        # Use our smart scoring system for better bullet points
        sentence_scores = score_sentences(sentences, cleaned_text)

        # Get top sentences by score
        top_sentences = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:max_points]

        # Sort by original order to maintain narrative flow
        top_sentences = sorted(top_sentences, key=lambda x: x[0])

        # Create bullet points with better formatting
        bullet_points = []
        for i, score in top_sentences:
            sentence = sentences[i].strip()
            # Clean up the sentence for bullet point format
            if sentence.endswith('.'):
                sentence = sentence[:-1]
            bullet_points.append(f"• {sentence}")

        return '\n'.join(bullet_points)
    except Exception as e:
        print(f"Error in bullet point summarization: {e}")
        return f"• {text}"

def extract_key_phrases(text, max_phrases=10):
    """Extract key phrases from the text"""
    try:
        cleaned_text = clean_text(text)
        words = nltk.word_tokenize(cleaned_text.lower())

        # Remove stopwords and short words
        stop_words = set(nltk.corpus.stopwords.words('english'))
        words = [word for word in words if word.isalpha() and len(word) > 3 and word not in stop_words]

        # Get word frequency
        word_freq = Counter(words)

        # Get most common words as key phrases
        key_phrases = [word.title() for word, _ in word_freq.most_common(max_phrases)]

        return ', '.join(key_phrases)
    except Exception as e:
        print(f"Error in key phrase extraction: {e}")
        return "Unable to extract key phrases"

def abstractive_summarize(text, target_length=3):
    """Create an abstractive summary by combining key concepts"""
    try:
        cleaned_text = clean_text(text)
        sentences = nltk.sent_tokenize(cleaned_text)

        if len(sentences) <= target_length:
            return cleaned_text

        # Extract key concepts and themes
        words = nltk.word_tokenize(cleaned_text.lower())
        stop_words = set(nltk.corpus.stopwords.words('english'))
        important_words = [word for word in words if word.isalpha() and len(word) > 3 and word not in stop_words]

        # Get most frequent important words
        word_freq = Counter(important_words)
        key_concepts = [word for word, _ in word_freq.most_common(10)]

        # Find sentences that contain multiple key concepts
        concept_sentences = []
        for sentence in sentences:
            sentence_words = set(nltk.word_tokenize(sentence.lower()))
            concept_count = sum(1 for concept in key_concepts if concept in sentence_words)
            if concept_count >= 2:  # Sentences with multiple key concepts
                concept_sentences.append((sentence, concept_count))

        # Sort by concept density and select best ones
        concept_sentences.sort(key=lambda x: x[1], reverse=True)

        # Take the best sentences up to target length
        selected_sentences = [sent for sent, _ in concept_sentences[:target_length]]

        if len(selected_sentences) < target_length:
            # Fill with highest scoring sentences from our smart algorithm
            remaining_sentences = [s for s in sentences if s not in selected_sentences]
            if remaining_sentences:
                sentence_scores = score_sentences(remaining_sentences, cleaned_text)
                additional_needed = target_length - len(selected_sentences)
                top_additional = sorted(sentence_scores.items(), key=lambda x: x[1], reverse=True)[:additional_needed]
                selected_sentences.extend([remaining_sentences[i] for i, _ in top_additional])

        # Reorder sentences to maintain narrative flow
        ordered_sentences = []
        for sentence in sentences:
            if sentence in selected_sentences:
                ordered_sentences.append(sentence)
                if len(ordered_sentences) >= target_length:
                    break

        return ' '.join(ordered_sentences)

    except Exception as e:
        print(f"Error in abstractive summarization: {e}")
        return extractive_summarize(text, 'smart', target_length)



def clean_ai_response(response_text):
    """Clean and format AI response to make it more human-readable"""
    if not response_text:
        return response_text

    # Remove common AI response patterns
    cleaned = response_text.strip()

    # Remove introductory phrases
    intro_patterns = [
        r"^Here's a concise summary.*?:\s*",
        r"^Here is a summary.*?:\s*",
        r"^Summary:\s*",
        r"^The summary is:\s*",
        r"^In summary:\s*",
        r"^To summarize:\s*",
        r"^Here's what the text covers:\s*",
        r"^The text discusses:\s*",
        r"^The main points are:\s*",
        r"^Key points:\s*"
    ]

    for pattern in intro_patterns:
        cleaned = re.sub(pattern, "", cleaned, flags=re.IGNORECASE | re.MULTILINE)

    # Remove concluding phrases
    conclusion_patterns = [
        r"\s*This summary captures.*$",
        r"\s*These are the main.*$",
        r"\s*This covers the essential.*$",
        r"\s*The summary includes.*$",
        r"\s*This provides.*$"
    ]

    for pattern in conclusion_patterns:
        cleaned = re.sub(pattern, "", cleaned, flags=re.IGNORECASE | re.MULTILINE)

    # Clean up numbered/bulleted lists to be more natural
    # Convert "1. **Point**: Description" to "Point: Description"
    cleaned = re.sub(r'^\d+\.\s*\*\*(.*?)\*\*:\s*', r'\1: ', cleaned, flags=re.MULTILINE)

    # Convert "1. Point" to "Point"
    cleaned = re.sub(r'^\d+\.\s*', '', cleaned, flags=re.MULTILINE)

    # Remove excessive markdown formatting
    cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', cleaned)  # Remove bold
    cleaned = re.sub(r'\*(.*?)\*', r'\1', cleaned)      # Remove italic

    # Clean up multiple spaces and newlines
    cleaned = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned)  # Max 2 consecutive newlines
    cleaned = re.sub(r' +', ' ', cleaned)                # Multiple spaces to single

    # If it's a list, convert to paragraph format
    lines = cleaned.split('\n')
    if len(lines) > 1 and all(line.strip() for line in lines):
        # Check if it looks like a list
        list_indicators = sum(1 for line in lines if re.match(r'^\s*[-•]\s*', line) or ':' in line)
        if list_indicators > len(lines) * 0.6:  # More than 60% are list items
            # Convert to paragraph
            processed_lines = []
            for line in lines:
                line = line.strip()
                if line:
                    # Remove list markers
                    line = re.sub(r'^[-•]\s*', '', line)
                    # Ensure proper sentence ending
                    if line and not line.endswith(('.', '!', '?')):
                        line += '.'
                    processed_lines.append(line)
            cleaned = ' '.join(processed_lines)

    # Final cleanup
    cleaned = cleaned.strip()

    # Ensure proper sentence structure
    if cleaned and not cleaned.endswith(('.', '!', '?')):
        cleaned += '.'

    return cleaned

def openrouter_summarize(text, length='medium'):
    """Summarize text using OpenRouter API with Mistral model"""
    try:
        # Check rate limits
        can_request, message = openrouter_rate_limiter.can_make_request()
        if not can_request:
            return None, message

        # Get API key from environment
        api_key = os.environ.get('OPENROUTER_API_KEY')
        if not api_key:
            return None, "OpenRouter API key not found. Please set OPENROUTER_API_KEY in .env file."

        # Prepare the request
        url = "https://openrouter.ai/api/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "HTTP-Referer": os.environ.get('SITE_URL', 'http://localhost:3000'),
            "X-Title": os.environ.get('SITE_NAME', 'Vision-Script Text Summarizer'),
            "Content-Type": "application/json"
        }

        # Create summarization prompt based on length
        if length == 'short':
            length_instruction = "in 2-3 sentences"
            max_tokens = 150
        elif length == 'long':
            length_instruction = "in 5-7 sentences"
            max_tokens = 400
        else:  # medium
            length_instruction = "in 3-4 sentences"
            max_tokens = 250

        # Truncate text if too long
        max_chars = 8000  # Conservative limit
        if len(text) > max_chars:
            text = text[:max_chars] + "..."

        prompt = f"""Summarize the following text {length_instruction}. Write in a natural, flowing style without introductory phrases, bullet points, or numbered lists. Focus on the main points:

{text}"""

        payload = {
            "model": "mistralai/mistral-small-3.2-24b-instruct:free",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "max_tokens": max_tokens,
            "temperature": 0.3
        }

        # Make the request
        response = requests.post(url, headers=headers, json=payload, timeout=60)

        if response.status_code == 200:
            result = response.json()
            if 'choices' in result and len(result['choices']) > 0:
                raw_summary = result['choices'][0]['message']['content'].strip()
                if raw_summary:
                    # Clean the AI response to make it more human-readable
                    cleaned_summary = clean_ai_response(raw_summary)
                    # Record successful request
                    openrouter_rate_limiter.record_request()
                    return cleaned_summary, "success"
                else:
                    return None, "No summary generated"
            else:
                return None, "Invalid response format"
        elif response.status_code == 401:
            return None, "Invalid API key. Please check your OPENROUTER_API_KEY in .env file."
        elif response.status_code == 429:
            # Rate limited by OpenRouter
            openrouter_rate_limiter.is_rate_limited = True
            openrouter_rate_limiter.rate_limit_until = datetime.now() + timedelta(minutes=5)
            return None, "Rate limit reached. Please try again in a few minutes."
        elif response.status_code == 402:
            return None, "Insufficient credits. Please check your OpenRouter account balance."
        else:
            return None, f"API error: {response.status_code} - {response.text}"

    except requests.exceptions.Timeout:
        return None, "Request timeout. Please try again."
    except requests.exceptions.RequestException as e:
        return None, f"Network error: {str(e)}"
    except Exception as e:
        return None, f"Error: {str(e)}"

@app.route('/openrouter_status', methods=['GET'])
def openrouter_status():
    """Get OpenRouter API rate limit status"""
    status = openrouter_rate_limiter.get_status()
    return jsonify(status)

@app.route('/test_openrouter', methods=['POST'])
def test_openrouter():
    """Test OpenRouter API with a simple request"""
    try:
        test_text = "This is a test text for summarization. It should be long enough to test the API properly."
        summary, status = openrouter_summarize(test_text, 'medium')
        return jsonify({
            "test_text": test_text,
            "summary": summary,
            "status": status,
            "success": summary is not None
        })
    except Exception as e:
        return jsonify({
            "error": str(e),
            "success": False
        }), 500

@app.route('/summarize_text', methods=['POST'])
def summarize_text():
    """Summarize text with various options"""
    try:
        data = request.get_json()

        if not data or 'text' not in data:
            return jsonify({"error": "No text provided"}), 400

        text = data['text'].strip()
        if not text:
            return jsonify({"error": "Empty text provided"}), 400

        # Get parameters
        algorithm = data.get('algorithm', 'textrank')  # textrank, lsa, luhn, abstractive
        smart_option = data.get('smart_option', 'local_smart')  # local_smart, huggingface
        summary_type = data.get('type', 'paragraph')  # paragraph, bullets, keyphrases
        length = data.get('length', 'medium')  # short, medium, long

        # Determine sentence count based on length and text size
        text_sentences = len(nltk.sent_tokenize(text))
        if length == 'short':
            sentences_count = max(1, min(3, text_sentences // 4))
        elif length == 'long':
            sentences_count = max(3, min(8, text_sentences // 2))
        else:  # medium
            sentences_count = max(2, min(5, text_sentences // 3))

        # Handle smart summarization first
        if smart_option == 'openrouter':
            # Use OpenRouter API for summarization
            or_summary, or_status = openrouter_summarize(text, length)

            if or_summary:
                if summary_type == 'bullets':
                    summary = bullet_point_summarize(or_summary, sentences_count)
                elif summary_type == 'keyphrases':
                    summary = extract_key_phrases(or_summary, sentences_count * 2)
                else:
                    summary = or_summary
            else:
                # Fallback to local smart if OpenRouter fails
                return jsonify({
                    "error": f"OpenRouter API error: {or_status}",
                    "fallback_available": True,
                    "status": "error"
                }), 400
        else:
            # Use local algorithms (including local smart)
            if summary_type == 'bullets':
                if algorithm == 'abstractive':
                    # Create abstractive summary first, then convert to bullets
                    para_summary = abstractive_summarize(text, sentences_count)
                    summary = bullet_point_summarize(para_summary, sentences_count)
                else:
                    summary = bullet_point_summarize(text, sentences_count)
            elif summary_type == 'keyphrases':
                summary = extract_key_phrases(text, sentences_count * 2)
            else:  # paragraph
                if algorithm == 'abstractive':
                    summary = abstractive_summarize(text, sentences_count)
                elif smart_option == 'local_smart':
                    summary = extractive_summarize(text, 'smart', sentences_count)
                else:
                    summary = extractive_summarize(text, algorithm, sentences_count)

        # Calculate statistics
        stats = get_text_statistics(text, summary)

        return jsonify({
            "original_text": text,
            "summary": summary,
            "algorithm": algorithm,
            "smart_option": smart_option,
            "type": summary_type,
            "length": length,
            "statistics": stats,
            "status": "success"
        })

    except Exception as e:
        print(f"Exception in summarize_text: {str(e)}")  # Debug log
        import traceback
        traceback.print_exc()  # Print full traceback
        return jsonify({
            "error": str(e),
            "status": "error"
        }), 500

@app.route('/extract_id_data', methods=['POST'])
def extract_id_data():
    file = request.files['image']
    lang_code = request.form.get('language', 'en').lower()
    
    file_path = os.path.join(tempfile.gettempdir(), file.filename)
    file.save(file_path)
    
    extracted_text = extract_text(file_path, lang_code=lang_code)
    extracted_data = clean_extracted_text(extracted_text, lang_code=lang_code)
    
    tmp_path = os.path.join(tempfile.gettempdir(), "id_card_data.xlsx")
    workbook = xlsxwriter.Workbook(tmp_path)
    worksheet = workbook.add_worksheet()
    
    # Add a header row with the detected language
    header_format = workbook.add_format({'bold': True, 'bg_color': '#333333', 'font_color': 'white'})
    language_name = SUPPORTED_LANGUAGES.get(lang_code, 'Unknown')
    worksheet.write(0, 0, 'Field', header_format)
    worksheet.write(0, 1, f'Value (Language: {language_name})', header_format)
    
    row = 1
    for key, value in extracted_data.items():
        worksheet.write(row, 0, key)
        worksheet.write(row, 1, value)
        row += 1
    
    workbook.close()
    
    return send_file(tmp_path, as_attachment=True, download_name="id_card_data.xlsx")

@app.route('/camera_feed', methods=['POST'])
def camera_feed():
    try:
        data = request.get_json()
        if not data or 'image' not in data:
            return jsonify({"error": "No image data provided"}), 400

        model = data.get('model', 'easyocr').lower()
        lang_code = data.get('language', 'en').lower()
        image_bytes = base64.b64decode(data['image'])
        nparr = np.frombuffer(image_bytes, np.uint8)
        frame = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        if model == 'pytesseract':
            # Convert to grayscale
            gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
            
            # Apply thresholding to remove noise
            gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
            
            # Resize image to improve text recognition
            scale_percent = 150  # Increase size by 150%
            width = int(gray.shape[1] * scale_percent / 100)
            height = int(gray.shape[0] * scale_percent / 100)
            gray = cv2.resize(gray, (width, height), interpolation=cv2.INTER_CUBIC)
            
            # Convert back to PIL image for Pytesseract
            img = Image.fromarray(gray)
            
            # Use optimized OCR settings with language support
            config = '--psm 6 --oem 3'
            if lang_code != 'en':
                extracted_text = pytesseract.image_to_string(img, lang=lang_code, config=config)
            else:
                extracted_text = pytesseract.image_to_string(img, config=config)
            
            return jsonify({
                "detections": [{
                    "text": extracted_text,
                    "bbox": {"x": 0, "y": 0, "width": 100, "height": 20},
                    "status": "success"
                }]
            })
        else:
            # EasyOCR implementation with language support
            reader = get_reader(lang_code)
            results = reader.readtext(frame)
            detections = []
            for (bbox, text, prob) in results:
                (tl, tr, br, bl) = bbox
                x = int(tl[0])
                y = int(tl[1])
                width = int(br[0] - tl[0])
                height = int(br[1] - tl[1])
                detections.append({
                    "text": text,
                    "bbox": {"x": x, "y": y, "width": width, "height": height},
                    "status": "success"
                })
            return jsonify({"detections": detections})

    except Exception as e:
        return jsonify({
            "error": str(e),
            "status": "error"
        }), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
